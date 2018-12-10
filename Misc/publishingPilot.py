import docx
import sys
import time
import os

def getText(y):
	fullText = []
	for para in y.paragraphs:
		fullText.append(para.text)
	return '\n'.join(fullText)

def constructNewDoc():
	oldDocFilePath = input('Enter filepath for an existing Word Document (.docx): ')
	newDoc = docx.Document(oldDocFilePath)
	bodyText = getText(newDoc)
	newDoc._body.clear_content()
	newDoc.add_heading('Test Heading')
	newDoc.add_heading('Test Subheading', level=2)
	newBody = newDoc.add_paragraph()
	runner = newBody.add_run(bodyText)
	runner.bold = True
	runner.italic = True
	font = runner.font
	font.color.rgb = docx.shared.RGBColor(0x42, 0x24, 0xE9)
	cwd = os.getcwd()
	newDoc.save('newDoc.docx')

	
if __name__ == '__main__':
	constructNewDoc()